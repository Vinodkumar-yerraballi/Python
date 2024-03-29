# -*- coding: utf-8 -*-
"""pdf_to_json.ipynb

Automatically generated by Colaboratory.

Original file is located at
    https://colab.research.google.com/drive/1DfzXfkfviATu8fBOB5QeC3BBHdDDWzF-
"""

!pip install PyPDF2

!pip install tabula_py

!pip install docx2txt

# Final function
import PyPDF2
import tabula
from tabula.io import read_pdf
import json
import os
import docx2txt


def pdf_to_json(pdf_path):
    # Open the PDF file
    # with open(pdf_path, 'rb') as file:
        # Create a PyPDF2 reader object
        pdf_reader = PyPDF2.PdfReader(pdf_path)
        # Check if the PDF has any pages
        # write the conditon if no pages return to none
        if pdf_reader.pages == 0:
            return None
        # Extract the table data from the first page of the PDF using tabula-py
        df = read_pdf(pdf_path)
        # Check if the page contains any tables
        if len(df) == 0:
            return None
        else:
            # Convert the table data to a list of JSON objects
            #create a empty list
            table = []
            headers = df[0].columns.tolist()
            # print(headers)
            for index, row in df[0].iterrows():
              #create a dictionary
                row_dict = {}
                # header in headers
                for header in headers:
                  #then diction heades covert to row header in string format
                    row_dict[header] = str(row[header])
                    #finally append to the dicion then dump the model
                table.append(row_dict)
            with open('data.json', 'w') as f:
                    json.dump(table, f) 
            return

table = pdf_to_json('/content/My Resume .pdf')
if table is not None:
    print(table)
# with open(table,'w') as file:
#       print(file)
# else:
#     print("The PDF file does not contain any tables.")



pdf_path=('/content/My Resume .pdf')
df = read_pdf(pdf_path, pages='all', multiple_tables=True,guess='area')
sliced=slice(0,-1,1)
element=df[:]

def pdf_to_json(file_path):
    with open(file_path,'rb') as file:
        reader=PyPDF2.PdfReader(file)
        if reader.pages==0:
            return None
        create_table=read_pdf(file_path,pages='all',multiple_tables=True)
        if len(create_table)==0:
            return None
        else:
            list_1=[]
            element=create_table[0].columns.tolist()
            for index,row in create_table[0].iterrows():
                    # print(index,row)
                    row_dict={}
                    for elements in element:
                      row_dict[elements]=str(row[elements])
                    list_1.append(row_dict)
            return json.dumps(list_1)
            # table=element.to_json(orient='index')
            # list_1.append(table)
table = pdf_to_json('/content/My Resume .pdf')
if table is not None:
    print(table)
# else:
#     print("The PDF file does not contain any tables.")

input_folder = 'C:/Users/Admin/Downloads/intern/input/'
output_folder = 'C:/Users/Admin/Downloads/intern/output'

if not os.path.exists(output_folder):
    os.makedirs(output_folder)

for root, dirs, files in os.walk(input_folder):
    for file in files:
        file_path = os.path.join(root, file)
        file_name, file_ext = os.path.splitext(file)
        if file_ext == '.pdf':
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ''
                for page in pdf_reader.pages:
                    text += page.extract_text()
        elif file_ext == '.docx':
            text = docx2txt.process(file_path)
        else:
            continue
        
        output_path = os.path.join(output_folder, file_name + '.txt')
        with open(output_path, 'w',encoding="utf-8") as output_file:
            output_file.write(text)

pip install python-docx

import docx
content = docx.Document('/content/VinodKumar Yerraballi.docx')
content

import json
import docx

# Load the .docx file
doc = docx.Document('/content/My Resume .docx')

# Create an empty list to store the table data
table_data = []

# Iterate over all tables in the document
for table in doc.tables:
    # Create an empty dictionary to store the table information
    table_dict = {}

    # Iterate over all rows in the table
    for i, row in enumerate(table.rows):
        # Create an empty list to store the row data
        row_data = []

        # Iterate over all cells in the row
        for cell in row.cells:
            # Append the text in the cell to the row data list
            row_data.append(cell.text)

        # Add the row data to the table dictionary
        table_dict[f"Row {i+1}"] = row_data

    # Add the table dictionary to the table data list
    table_data.append(table_dict)

# Convert the table data list to JSON format
table_json = json.dumps(table_data)

# Print the JSON data
print(table_json)

from tabulate import tabulate

import json
import tabula

# Load the PDF file and extract the table data
tables = tabula.read_pdf('/content/My Resume .pdf')

# Create an empty list to store the table data
table_data = []

# Iterate over all tables in the document
for table in tables:
    # print(table)
    # Create an empty dictionary to store the table information
    table_dict = {}

    # Iterate over all rows in the table
    for i, row in table.iterrows():
        # Create an empty list to store the row data
        row_data = []

        # Iterate over all cells in the row
        for cell in row:
            # Append the text in the cell to the row data list
            row_data.append(cell)

        # Add the row data to the table dictionary
        table_dict[f"Row {i}"] = row_data

    # Add the table dictionary to the table data list
    table_data.append(table_dict)

with open(f'My Resume.json','w') as file:
  json.dump(table_data,file)

tables = tabula.read_pdf('/content/excellent-resume-format-free.pdf', pages='all',multiple_tables=True,)
tables

# save them in a folder
folder_name = "tables"
if not os.path.isdir(folder_name):
    os.mkdir(folder_name)
# iterate over extracted tables and export as excel individually
for i, table in enumerate(tables, start=1):
    table.to_json(os.path.join(folder_name, f"file_name{i}.json"), index='table')

tabula.convert_into("/content/My Resume .pdf", "output.json", output_format="json", pages="all")

pip install camelot-py

!pip install pdf2docx

from pdf2docx import Converter
import os

# # # dir_path for input reading and output files & a for loop # # #

path_input = '/content/My Resume .pdf'
path_output = '/content/tables'

for file in os.listdir(path_input):
    cv = Converter(path_input+file)
    cv.convert(path_output+file+'.docx', start=0, end=None)
    cv.close()
    print(file)

import ctypes
from ctypes.util import find_library
find_library("".join(("gsdll", str(ctypes.sizeof(ctypes.c_voidp) * 8), ".dll")))

from tabulate import tabulate
from tabula import read_pdf
tables =read_pdf('/content/My Resume .pdf')
tables



import docx2txt
def docx_to_json(file_path):
    with open(file_path, 'rb') as doc_file:
      doc = docx.Document(doc_file)
      table_data = []
      for table in doc.tables:
        table_dict = {}
        for i, row in enumerate(table.rows):
          row_data = []
          for cell in row.cells:
            table_dict[i] = row_data
            table_data.append(table_dict)
            with open(f'{file_path}.json','w') as file:
                       json.dump(table_data,file)
            if doc not in table_data:
                            text =' '
                            text=docx2txt.process(file_path)
                            with open(f'{file_path}.txt','w') as file:
                                  file.write(text)
        return

table=docx_to_json('/content/My Resume .docx')
table